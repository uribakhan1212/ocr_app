"""
Document generation utilities for creating Word documents
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from PIL import Image
import tempfile
import os
from typing import List, Dict, Optional
import pandas as pd
from datetime import datetime

class WordDocumentGenerator:
    """Generate Word documents from extracted content"""
    
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """Setup custom styles for the document"""
        # Create custom styles
        styles = self.doc.styles
        
        # Title style
        if 'Custom Title' not in [style.name for style in styles]:
            title_style = styles.add_style('Custom Title', 1)  # 1 = paragraph style
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.size = Pt(18)
            title_font.bold = True
            title_font.color.rgb = RGBColor(0, 51, 102)
        
        # Heading style
        if 'Custom Heading' not in [style.name for style in styles]:
            heading_style = styles.add_style('Custom Heading', 1)
            heading_font = heading_style.font
            heading_font.name = 'Calibri'
            heading_font.size = Pt(14)
            heading_font.bold = True
            heading_font.color.rgb = RGBColor(0, 51, 102)
        
        # Body text style
        if 'Custom Body' not in [style.name for style in styles]:
            body_style = styles.add_style('Custom Body', 1)
            body_font = body_style.font
            body_font.name = 'Calibri'
            body_font.size = Pt(11)
    
    def add_header(self, title: str = "Extracted Document Content"):
        """Add document header with title"""
        # Add title
        title_para = self.doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add generation timestamp
        timestamp = datetime.now().strftime("%B %d, %Y at %I:%M %p")
        timestamp_para = self.doc.add_paragraph(f"Generated on {timestamp}")
        timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add separator line
        self.doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph()  # Empty line
    
    def add_original_image(self, image: Image.Image, caption: str = "Original Image"):
        """Add the original image to the document"""
        if image is None:
            return
        
        # Add section heading
        self.doc.add_heading(caption, level=1)
        
        # Save image temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
            # Resize image if too large
            max_width = 6.5  # inches
            img_width, img_height = image.size
            
            if img_width > img_height:
                width_inches = min(max_width, 6)
                height_inches = (img_height * width_inches) / img_width
            else:
                height_inches = min(max_width, 8)
                width_inches = (img_width * height_inches) / img_height
            
            image.save(tmp_file.name, 'PNG')
            
            # Add image to document
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.add_picture(tmp_file.name, width=Inches(width_inches))
            
            # Clean up temp file
            os.unlink(tmp_file.name)
        
        self.doc.add_paragraph()  # Empty line after image
    
    def add_extracted_text(self, text_content: str, preserve_formatting: bool = True):
        """Add extracted text content to the document"""
        if not text_content.strip():
            return
        
        # Add section heading
        self.doc.add_heading('Extracted Text Content', level=1)
        
        if preserve_formatting:
            # Split into paragraphs and maintain structure
            paragraphs = text_content.split('\n\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    # Check if it looks like a heading (short, all caps, etc.)
                    if self._is_likely_heading(para_text.strip()):
                        heading_para = self.doc.add_heading(para_text.strip(), level=2)
                    else:
                        para = self.doc.add_paragraph(para_text.strip())
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            # Add as single block
            para = self.doc.add_paragraph(text_content.strip())
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.doc.add_paragraph()  # Empty line after text
    
    def add_tables(self, tables: List[List[str]], table_titles: Optional[List[str]] = None):
        """Add detected tables to the document"""
        if not tables:
            return
        
        # Add section heading
        self.doc.add_heading('Detected Tables', level=1)
        
        for i, table_data in enumerate(tables):
            if not table_data:
                continue
            
            # Add table title if provided
            if table_titles and i < len(table_titles):
                self.doc.add_heading(table_titles[i], level=2)
            else:
                self.doc.add_heading(f'Table {i + 1}', level=2)
            
            # Determine table dimensions
            max_cols = max(len(row) for row in table_data) if table_data else 0
            
            if max_cols > 0:
                # Create table
                table = self.doc.add_table(rows=len(table_data), cols=max_cols)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Fill table data
                for row_idx, row_data in enumerate(table_data):
                    row_cells = table.rows[row_idx].cells
                    
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < len(row_cells):
                            cell = row_cells[col_idx]
                            cell.text = str(cell_data).strip()
                            
                            # Make first row bold (header)
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
                
                self.doc.add_paragraph()  # Empty line after table
    
    def add_processing_summary(self, stats: Dict):
        """Add processing summary and statistics"""
        self.doc.add_heading('Processing Summary', level=1)
        
        # Create summary table
        summary_table = self.doc.add_table(rows=len(stats), cols=2)
        summary_table.style = 'Table Grid'
        
        for i, (key, value) in enumerate(stats.items()):
            row_cells = summary_table.rows[i].cells
            row_cells[0].text = str(key)
            row_cells[1].text = str(value)
            
            # Make first column bold
            for paragraph in row_cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
    
    def add_confidence_scores(self, text_blocks: List[Dict]):
        """Add confidence scores for extracted text"""
        if not text_blocks:
            return
        
        self.doc.add_heading('Text Extraction Confidence', level=1)
        
        # Calculate average confidence
        avg_confidence = sum(block['confidence'] for block in text_blocks) / len(text_blocks)
        
        para = self.doc.add_paragraph(f"Average confidence score: {avg_confidence:.1%}")
        
        # Add confidence breakdown
        high_conf = sum(1 for block in text_blocks if block['confidence'] > 0.8)
        medium_conf = sum(1 for block in text_blocks if 0.5 <= block['confidence'] <= 0.8)
        low_conf = sum(1 for block in text_blocks if block['confidence'] < 0.5)
        
        conf_table = self.doc.add_table(rows=4, cols=2)
        conf_table.style = 'Table Grid'
        
        conf_data = [
            ('Confidence Level', 'Text Blocks'),
            ('High (>80%)', str(high_conf)),
            ('Medium (50-80%)', str(medium_conf)),
            ('Low (<50%)', str(low_conf))
        ]
        
        for i, (label, value) in enumerate(conf_data):
            row_cells = conf_table.rows[i].cells
            row_cells[0].text = label
            row_cells[1].text = value
            
            if i == 0:  # Header row
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
    
    def _is_likely_heading(self, text: str) -> bool:
        """Determine if text is likely a heading"""
        # Simple heuristics for heading detection
        if len(text) > 100:  # Too long to be a heading
            return False
        
        if text.isupper() and len(text) < 50:  # All caps, short
            return True
        
        if text.endswith(':') and len(text) < 50:  # Ends with colon
            return True
        
        # Check for common heading patterns
        heading_patterns = ['chapter', 'section', 'part', 'introduction', 'conclusion']
        if any(pattern in text.lower() for pattern in heading_patterns):
            return True
        
        return False
    
    def save_document(self) -> bytes:
        """Save document and return as bytes"""
        import io
        doc_buffer = io.BytesIO()
        self.doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer.getvalue()
    
    def create_complete_document(self, 
                                text_content: str,
                                tables: List[List[str]],
                                original_image: Optional[Image.Image] = None,
                                text_blocks: Optional[List[Dict]] = None,
                                include_stats: bool = True) -> bytes:
        """Create a complete Word document with all content"""
        
        # Add header
        self.add_header()
        
        # Add original image if provided
        if original_image:
            self.add_original_image(original_image)
        
        # Add extracted text
        if text_content:
            self.add_extracted_text(text_content)
        
        # Add tables
        if tables:
            self.add_tables(tables)
        
        # Add processing summary
        if include_stats and text_blocks:
            stats = {
                'Total Text Blocks': len(text_blocks),
                'Total Words': len(text_content.split()) if text_content else 0,
                'Tables Detected': len(tables) if tables else 0,
                'Processing Engine': 'EasyOCR + Tesseract'
            }
            self.add_processing_summary(stats)
            
            # Add confidence scores
            self.add_confidence_scores(text_blocks)
        
        return self.save_document()